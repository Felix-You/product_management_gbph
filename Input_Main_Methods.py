import os,threading,math,time,json, sqlite3,copy
import ConnSqlite as CS
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from MeetingModeDialogUi import Ui_Dialog
from PyQt5.QtWidgets import *
from PyQt5 import QtWidgets,QtGui
from PyQt5.QtCore import *
from ID_Generate import Snow
from FilePathInit import userDir, workingDir
import csv, xlwt,openpyxl, DataView, datetime,types, RedefinedWidget, GColour, DataCenter
from collections import OrderedDict
from DataView import DF_Ratio, FIX_SIZE_WIDGET_SCALING
from RedefinedWidget import ComboBoxDialog
from functools import partial

from core.KitModels import Person, PersonLog


class QComboBox(QComboBox):
    def wheelEvent(self, QWheelEvent):
        QWheelEvent.ignore()

def new_focusOutEvent(self,e):
    if e.reason() == Qt.ActiveWindowFocusReason:
        return
    else:
        QLineEdit.focusOutEvent(self, e)

class MeetingRecordDialog(QtWidgets.QDialog, Ui_Dialog):
    def task_edit_focusOutEvent(self,cls, e):
        text = cls.toPlainText().strip()
        if text:
            self.tableWidget.cellWidget(self.tableWidget.currentRow(),self.col_name_index['待解决']).setEnabled(True)
            self.tableWidget.cellWidget(self.tableWidget.currentRow(),self.col_name_index['添加追踪']).setEnabled(True)
        else:
            self.tableWidget.cellWidget(self.tableWidget.currentRow(),self.col_name_index['待解决']).setEnabled(False)
            self.tableWidget.cellWidget(self.tableWidget.currentRow(),self.col_name_index['添加追踪']).setEnabled(False)
            self.tableWidget.cellWidget(self.tableWidget.currentRow(),self.col_name_index['添加追踪']).setChecked(False)
        QTextEdit.focusOutEvent(cls, e)

    def newResizeEvent(self,cls, a0: QtGui.QResizeEvent) -> None:
        self.lineEdit_2.setFixedWidth(cls.geometry().width())
        QTextEdit.resizeEvent(cls, a0)

    def __init__(self, client_name:str,client_id:str, company_meeting_id:str = None, parent=None):
        super(MeetingRecordDialog, self).__init__()
        self.setupUi(self)
        DataView.ResAdaptor.init_ui_size(self)
        self.setAttribute(Qt.WA_DeleteOnClose)
        self.setWindowTitle(client_name + '-会议纪要')
        # self.setAttribute(Qt.WA_Hover)
        self.parent = parent
        self.filepath_model = userDir

        # self.setStyleSheet('QTextEdit {'
        #                            'border-width: 1px;'
        #                            'border-style: solid;'
        #                            'border-color: qlineargradient(spread:pad, x1:0.5, y1:1, x2:0.5, y2:0, '
        #                            'stop:0 rgba(0, 113, 255, 255), stop:1 rgba(91, 171, 252, 255));}')
        self.initUserTheme(self)
        self.tableWidget.setStyleSheet('QTextEdit{border-width: 0px;background-color: rgba(255, 255, 255,0);}'
                                       "QComboBox{border-width: 0px;background-color: rgba(255, 255, 255,0);}"
                                       "QComboBox::drop-down {subcontrol-origin: padding;subcontrol-position: top right;"
                                                           "width: 15px;border-left-width: 0px;"
                                                           "border-top-right-radius: 3px;border-bottom-right-radius: 3px;}"
                                        "QTableWidget::QTextEdit{border-width: 0px}"
                                        "QTableWidget::item{border-width: 1px; border-style}"
                                       "QScrollBar:vertical {"
                                                                         "width: 4px;"
                                                                   "background: transparent;"
                                                                   "border: 0px;"
                                                                   "margin: 0px 0px 0px 0px;"
                                                                   "}"
                                                                   "QScrollBar::handle:vertical"
                                                                   "{"
                                                                   "background:silver;"
                                                                   "border-radius:2px;"
                                                                   "}"
                                                                   "QScrollBar::handle:vertical:hover"
                                                                   "{"
                                                                   "background:lightskyblue;"
                                                                   "border-radius:2px;"
                                                                   "}"
                                                                   "QScrollBar::add-line:vertical"
                                                                   "{"
                                                                   "background: transparent;"
                                                                   "margin: 0px;"
                                                                   "border-width: 0px;"
                                                                   "height:0px;width:0px;"
                                                                   "subcontrol-position:bottom;"
                                                                   "subcontrol-origin: margin;"
                                                                   "}"
                                                                   "QScrollBar::sub-line:vertical"
                                                                   "{"
                                                                    "background: transparent;"
                                                                   "margin: 0px;"
                                                                   "border-width: 0px;"
                                                                   "height:0px;width:0px;"
                                                                   "subcontrol-position:top;"
                                                                   "subcontrol-origin: margin;"
                                                                   "}")
        self.pushButton.setMinimumSize(70, 30)

        self.client = client_name
        self.client_id = client_id
        self.company_meeting_id = company_meeting_id
        self.client_log_id =  Snow('clg').get()
        self.IDget = Snow('a')
        self.task_ID_get = Snow('ts')
        self.memo_id_get = Snow('mm')
        self.meeting_id_get = Snow('mt')
        self.status_id_get = Snow('st')
        self.todo_id_get = Snow('td')
        self.person_in_meeting = []
        self.log_id_track = OrderedDict()
        # 目的：在一次会议纪要编写窗口关闭前，所有的保存操作均可撤销，并以关闭窗口前最后一次保存的结果为最终结果
        # 初始化：若为已有的project，加载project原有记录中的参数，is_new_project为False；或初始化新参数，并标记为is_new_project为True；
        # 首次保存：全部标记为已保存；
        # 删除：将初始化的参数加入到已删除队列；若is_new_project，则需要先检查项目名称是否已在被删除的队列中；否则根据project_id来检查
        # 删除后重新保存：用track中的id删除追加型记录，或者覆盖覆写类的的记录；其中proj_list和proj_status_log通常为覆写型，
        #             但若is_new_project为True则视为追加型记录
        #
        self.log_id_track_deleted = OrderedDict()
        #从数据文件加载窗口构造基本数据信息
        self.table_header = ['项目名称','详情记录','管线进度','机会情况','待办程度','子任务','待解决','添加追踪','项目备注','明确机会',
                             '预期订单','上线跟进','高亮关注','客户名称','项目ID']#表格标题字段
        self.table_header_key_new = ['product','meeting_desc','current_progress','status_code','is_critical','task_desc',
                                     'officejob_type','todo','memo_desc','clear_chance','order_tobe', 'in_act',
                                     'highlight','client', '_id']#新数据库字段
        #
        self.table_header_index = list(range(len(self.table_header)))
        header_index_worker = zip(self.table_header, self.table_header_index)
        # print('header_index_worker',header_index_worker)
        self.col_name_index = dict(header_index_worker)  #字典{表头名:列数}
        header_key_worker = zip(self.table_header, self.table_header_key_new)
        self.col_name_key = dict(header_key_worker)#字典{表头名:数据库key_old}
        self.office_job_dict = DataCenter.office_job_dict
        self.secondary_editing = False#状态机，是否为二次编辑
        # print(self.col_name_index)
        # print(self.table_header)
        #获取管线进度选项
        progress_codes = sorted(DataCenter.progress_code.values())
        self.progress_items = [ DataCenter.progress_text[DataCenter.progress_code_r[i]] for i in progress_codes]
        self.comboBox_progress.addItems(self.progress_items)
        self.comboBox_progress.setCurrentIndex(0)
        #适配分辨率
        self.lineEdit.setMaximumSize(400*DF_Ratio,35*DF_Ratio)
        self.dateEdit.setFixedSize(150*DF_Ratio,25*DF_Ratio)
        self.textEdit_5.setMaximumSize(400*DF_Ratio,400*DF_Ratio)

        self.label_12.setMinimumWidth(75*DF_Ratio)
        self.label_15.setMinimumWidth(75*DF_Ratio)
        self.label_14.setMinimumWidth(75*DF_Ratio)
        self.label_13.setMinimumWidth(75*DF_Ratio)
        self.lineEdit_2.setMaximumSize(500*DF_Ratio,35*DF_Ratio)
        self.textEdit_3.setMaximumSize(500*DF_Ratio,400*DF_Ratio)
        self.textEdit_4.setMaximumSize(500*DF_Ratio,100*DF_Ratio)
        self.textEdit_6.setMaximumSize(500*DF_Ratio,100*DF_Ratio)
        self.textEdit_3.resizeEvent = types.MethodType(self.newResizeEvent, self.textEdit_3)

        self.comboBox_personnel.setMaximumSize(300*DF_Ratio ,30*DF_Ratio)
        self.textEdit.setMaximumSize(300*DF_Ratio, 90*DF_Ratio)
        self.textEdit_2.setMaximumSize(300*DF_Ratio, 65*DF_Ratio)

        self.pushButton_4.setMaximumSize(80*DF_Ratio,40*DF_Ratio)
        self.pushButton_2.setMaximumSize(80*DF_Ratio,40*DF_Ratio)
        self.pushButton_5.setMaximumSize(80*DF_Ratio,40*DF_Ratio)
        self.pushButton.setMaximumSize(80*DF_Ratio,40*DF_Ratio)
        self.comboBox_progress.setMinimumWidth(60*DF_Ratio)

        #建立tableWidget
        self.cols_count = len(self.table_header)
        self.tableWidget.setColumnCount(self.cols_count)
        self.tableWidget.setHorizontalHeaderLabels(self.table_header)#？？更改Header的来源
        self.tableWidget.setColumnWidth(1, 300*DF_Ratio)
        self.tableWidget.setColumnWidth(2, 70*DF_Ratio)#管线进度
        self.tableWidget.setColumnWidth(3, 140*DF_Ratio)#机会
        self.tableWidget.setColumnWidth(4, 75*DF_Ratio)#待办程度
        self.tableWidget.setColumnWidth(5, 200*DF_Ratio)#子任务
        self.tableWidget.setColumnWidth(6, 95*DF_Ratio)#待解决
        self.tableWidget.setColumnWidth(7, 70*DF_Ratio)#添加追踪
        self.tableWidget.setColumnWidth(8, 195*DF_Ratio)#项目备注
        self.tableWidget.setColumnWidth(9, 70*DF_Ratio)#明确机会
        self.tableWidget.setColumnWidth(10, 70*DF_Ratio)#预期订单
        self.tableWidget.setColumnWidth(11, 70*DF_Ratio)#上线跟进
        self.tableWidget.setColumnWidth(12, 70*DF_Ratio)#高亮关注
        self.tableWidget.setColumnWidth(13, 70*DF_Ratio)#客户名称
        self.tableWidget.setColumnWidth(14, 60*DF_Ratio)#项目ID
        # self.tableWidget.verticalHeader().setDefaultSectionSize(50)
        #日期编辑器
        self.dateEdit.setDateTime(QDateTime.currentDateTime())
        #按钮信号
        self.pushButton_5.clicked.connect(self.save_word_file)
        self.pushButton_6.clicked.connect(self.add_line)
        self.pushButton_4.clicked.connect(self.del_line)
        self.pushButton_2.clicked.connect(self.save_excel_file)
        # self.pushButton_8.clicked.connect(self.load_from_excel)
        self.pushButton.clicked.connect(self.update_dynamic)
        self.pushButton_up.clicked.connect(self.on_upButton_clicked)
        self.pushButton_down.clicked.connect(self.on_downButton_clicked)
        self.lineEdit_2.editingFinished.connect(self.is_project_exist)
        self.lineEdit_2.focusOutEvent = types.MethodType(new_focusOutEvent, self.lineEdit_2)

        #载入特殊控件
        self.setSpecialWidget()
        # #客户名称completer
        # customer_list = CS.getLinesFromTable('clients', columns_required=['short_name'],conditions={})
        # customer_list.pop()
        # self.customer_list = [item[0] for item in customer_list]
        # self.comp_model_1 = QStringListModel(self.customer_list)
        # self.m_completer = QCompleter(self.comp_model_1, self)
        # self.m_completer.setFilterMode(Qt.MatchContains)
        # self.lineEdit.setCompleter(self.m_completer)
        #项目名称completer
        self.set_project_completer()
        self.timer = QTimer(self)
        self.timerId = self.timer.timerId()
        self.timer.timeout.connect(self.auto_save_meeting_info)
        self.timer.start(10000)
        #载入已有会议信息
        if self.company_meeting_id:
            self.load_from_meeting_log()
        #设定客户名称输入
        self.lineEdit.setReadOnly(True)
        self.lineEdit.setText(self.client)

    def reject(self) -> None:
        pass

    def initUserTheme(self,cls:QWidget):
        theme = eval(self.filepath_model.getUserTheme())
        theme_dict = {0:'Aqua',1:'MacOS',2:'Ubuntu',3:'MaterialDark'}
        theme_name = theme_dict[theme]
        with open('./QSS/%s.qss'%theme_name) as qssfile:
            qss = qssfile.readlines()
            qss = ''.join(qss).strip('\n')
        cls.setStyleSheet(str(qss))

    def setSpecialWidget(self):
        #todo_check
        self.set_todo_checkBox = RedefinedWidget.SliderButton('添加追踪',self,colorChecked=GColour.TaskColour.TaskToDo,
                                                              inner_alpha=255,outer_checked_alpha=200,outer_unchecked_alpha=200)
        self.set_todo_checkBox.setFixedSize(QSize(60*DF_Ratio,20*DF_Ratio))
        self.set_todo_checkBox.setChecked(True)
        self.verticalLayout.replaceWidget(self.checkBox_todo,self.set_todo_checkBox)
        self.checkBox_todo.deleteLater()
        #comboBox_todo
        self.checkCombo_personnel = RedefinedWidget.CheckableComboBox(parent=self,init_all_checked=False)
        self.checkCombo_personnel.setMaximumWidth(250*DF_Ratio)
        self.formLayout_2.replaceWidget(self.comboBox_personnel,self.checkCombo_personnel)
        #特殊信号
        self.checkCombo_personnel.checkStatusChanged.connect(self.on_person_in_meeting)#防止每一次更新combbox的时候，都重新设置里一次连接
        self.pushButton_personnel.clicked.connect(self.on_editing_personnel)

        self.comboBox_personnel.deleteLater()
        self.initPersonnelSet()

    def initPersonnelSet(self):
        personnel = CS.getLinesFromTable('clients',{'_id':self.client_id},['_id','personnel'])
        personnel.pop()
        if personnel:
            personnel_json = personnel[0][1]
            self.personnel_data = json.loads(personnel_json) if personnel_json else []
        else:
            self.personnel_data = []
        self.set_comboBox_personnel()

    def set_comboBox_personnel(self,  is_renew = False):
        person_names = []
        person_ids = []
        for item in self.personnel_data:#name,title,telephone
            person_names.append(item['name'])
            person_ids.append(item['_id'])
        # if not is_renew:
        #     self.slot = lambda X: self.on_person_in_meeting(X, person_ids)
        # if is_renew:#重载comboBox之前先断开之前连接的槽函数，然后更新槽函数，完成重载之后重新再连接
        #     self.checkCombo_personnel.currentIndexChanged.disconnect(self.slot)
        #     self.slot = lambda X: self.on_person_in_meeting(X, person_ids)
        self.checkCombo_personnel.clear()
        self.checkCombo_personnel.addItems(person_names)
        self.textEdit.clear()

    def init_comboBox_personnel_checked(self, visited_people_json):
        #客户人员
        try:
            visited_people_data = json.loads(visited_people_json)
            self.person_in_meeting = visited_people_data
            person_in_meeting = []
            person_id_in_meeting = []
            for i, personnel_data in enumerate(visited_people_data):
                person_in_meeting.append(personnel_data['name'] + ' ' + personnel_data['tittle'])
                person_id_in_meeting.append(personnel_data['_id'])
            person_in_meeting_text = '\n'.join(person_in_meeting)
            #客户人员复选菜单
            person_check_status = []
            for i, person in enumerate(self.personnel_data):
                person_check_status.append(person['_id'] in person_id_in_meeting)
            self.checkCombo_personnel.setCheckStatus(person_check_status)
        except:
            person_in_meeting_text = visited_people_json

        self.textEdit.setText(person_in_meeting_text)

    def set_project_completer(self):
        project_list = CS.getLinesFromTable('proj_list',conditions={'client_id':self.client_id},columns_required=['product'])
        project_list.pop()
        self.project_list = [item[0] for item in project_list]
        self.comp_model_2 = QStringListModel(self.project_list)
        self.proj_completer = QCompleter(self.comp_model_2, self)
        self.proj_completer.setFilterMode(Qt.MatchContains)
        self.lineEdit_2.setCompleter(self.proj_completer)
        self.chance_status_filter = DataView.ChanceStatusFilterView()

    def nan_BoolJudge(self, val):#判断数据的布尔值
        try:
            if math.isnan(val):
                return False
            else:
                if val in ['FALSE', 'False', 'false', '0', '0.0']:
                    return False
                else:
                    return bool(val)
        except:
            if val in ['FALSE','False','false','0','0.0']:
                return False
            else:
                return bool(val)

    def is_number(self,st):
        try:
            float(st)
            return True
        except ValueError:
            pass
        try:
            import unicodedata
            unicodedata.numeric(st)
            return True
        except (TypeError, ValueError):
            pass
        return False

    def is_in_customer_list(self):
        txt = self.lineEdit.text()
        if txt == '':
            return
        elif txt in self.customer_list:
            return
        else:
            set_new_customer = QMessageBox.warning(self, "注意", "无记录的客户，是否新增客户？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if set_new_customer == QMessageBox.No:
                self.lineEdit.clear()
                self.lineEdit.setFocus()
            else:
                #todo:建立新客户信息
                return

    def is_project_exist(self):
        #判断该项目是否已经添加到列表，
        for i, key in enumerate(self.log_id_track.keys()):
            if self.log_id_track[key]['project_name'] == self.lineEdit_2.text():
                QMessageBox.about(self, '错误','请勿重复添加相同项目！')
                self.lineEdit_2.clear()
                return
        #判断输入的项目名称是否已经存在记录
        self.is_in_project_list()

    def is_in_project_list(self):
        '''判断输入的项目名称是否已经存在'''
        # if self.lineEdit_2.hasFocus():
        #     return
        txt = self.lineEdit_2.text()
        if txt == '':
            return
        if txt in self.project_list:
            #读取管线进度
            progress = CS.getLinesFromTable(table_name='proj_list', conditions={'product':txt, 'client_id':self.client_id},
                                            columns_required=['current_progress'])
            progress.pop()
            if progress:
                progress_xx = progress[0][0]
                self.comboBox_progress.setCurrentIndex(DataCenter.progress_code[progress_xx])
            return
        else:
            set_new_project = QMessageBox.question(self, "项目名称", "该客户无此项目记录，是否新增项目？",
                                                  QMessageBox.Yes | QMessageBox.No,
                                                   QMessageBox.No)
            if set_new_project == QMessageBox.No:
                self.lineEdit_2.setFocus()
                self.lineEdit_2.clear()
                # self.pushButton_6.clearFocus()
            else:
                #todo：建立新项目信息，绑定相应产品
                return

    def add_line(self):
        '''生成新的行'''
        if self.lineEdit_2.text() == '' or self.textEdit_3.toPlainText().strip() == "":
            return
        row_count = self.tableWidget.rowCount()
        #添加新的空行
        self.add_line_empty()
        #对空行进行初始化
        proj_name = self.lineEdit_2.text()
        client_id = self.client_id
        self.init_new_empty_line(proj_name,client_id)
        #填入新的记录
        #详情记录
        self.tableWidget.cellWidget(row_count,self.col_name_index['详情记录']).setText(str(self.textEdit_3.toPlainText()))
        #子任务
        self.tableWidget.cellWidget(row_count,self.col_name_index['子任务']).setText(str(self.textEdit_4.toPlainText()))
        #项目备注
        self.tableWidget.cellWidget(row_count,self.col_name_index['项目备注']).\
                setText(self.textEdit_6.toPlainText())
        self.set_todo_checkBox.setChecked(True)
        self.comboBox_progress.setCurrentIndex(0)
        self.lineEdit_2.clear()
        self.textEdit_3.clear()
        self.textEdit_4.clear()
        self.textEdit_6.clear()
        self.lineEdit_2.setFocus()

    def add_line_empty(self, nRow = None):
        '''生成新的空行'''
        if nRow is None:
            row_count = self.tableWidget.rowCount()
        else:
            row_count = nRow
        self.tableWidget.insertRow(row_count)
        #初始化整行
        #标签下拉菜单单独设置
        combo = QComboBox()
        combo_list_items = list(self.chance_status_filter.code_name_dict.items())
        combo_list_items.sort(key =lambda item: item[0])
        combo_list = [item[1] for item in combo_list_items]
        combo.addItems(combo_list)
        # combo.setStyleSheet("background-color: rgb(255, 255, 255);")
        self.tableWidget.setCellWidget(row_count , self.col_name_index['机会情况'] , combo)
        #管线进度下拉菜单
        current_progress_combo = QComboBox()
        self.tableWidget.setCellWidget(row_count , self.col_name_index['管线进度' ] , current_progress_combo)
        current_progress_combo.addItems(self.progress_items)

        #程度下拉菜单单独设置
        urgence_combo = QComboBox()
        urgence_combo.addItems([ '1.常规', '2.紧急' ])
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '待办程度' ] , urgence_combo)
        #明确机会CheckBox单独设置
        todo_checkBox = RedefinedWidget.SliderButton(colorChecked=GColour.TaskColour.TaskToDo,
                                                             outer_checked_alpha=205,
                                                             outer_unchecked_alpha=255,inner_alpha=205)
        todo_checkBox.setFixedSize(QSize(60*FIX_SIZE_WIDGET_SCALING, 30*FIX_SIZE_WIDGET_SCALING))
        self.tableWidget.setCellWidget(row_count , self.col_name_index['添加追踪'] , todo_checkBox)

        #明确机会CheckBox单独设置
        clear_chance_checkBox = RedefinedWidget.SliderButton(colorChecked=GColour.ProjectRGBColour.ProjectClearChance,
                                                             outer_checked_alpha=205,
                                                             outer_unchecked_alpha=255,inner_alpha=205)
        clear_chance_checkBox.setFixedSize(QSize(60*FIX_SIZE_WIDGET_SCALING,30*FIX_SIZE_WIDGET_SCALING))
        # clear_chance_checkBox.setColourChecked()
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '明确机会' ] , clear_chance_checkBox)
        #预期订单
        order_to_be_checkBox = RedefinedWidget.SliderButton(colorChecked=GColour.ProjectRGBColour.ProjectOrderTobe,
                                                            outer_checked_alpha=205,
                                                             outer_unchecked_alpha=255,inner_alpha=205)
        order_to_be_checkBox.setFixedSize(QSize(60*FIX_SIZE_WIDGET_SCALING,30*FIX_SIZE_WIDGET_SCALING))
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '预期订单' ] , order_to_be_checkBox)
        #上线跟进
        field_work_checkBox = RedefinedWidget.SliderButton(colorChecked=GColour.ProjectRGBColour.ProjectInAct,
                                                           outer_checked_alpha=205,
                                                             outer_unchecked_alpha=255,inner_alpha=205)
        field_work_checkBox.setFixedSize(QSize(60*FIX_SIZE_WIDGET_SCALING,30*FIX_SIZE_WIDGET_SCALING))
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '上线跟进' ] , field_work_checkBox)
        #待解决
        office_job_comboBox = QComboBox()
        items = self.office_job_dict.values()
        office_job_comboBox.addItems(items)
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '待解决' ] , office_job_comboBox)
        #高亮关注
        temp_handled_checkBox = RedefinedWidget.SliderButton(colorChecked=GColour.ProjectRGBColour.ProjectHighlight,
                                                             outer_checked_alpha=205,
                                                             outer_unchecked_alpha=255,inner_alpha=205)
        temp_handled_checkBox.setFixedSize(QSize(60*FIX_SIZE_WIDGET_SCALING,30*FIX_SIZE_WIDGET_SCALING))
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '高亮关注' ] , temp_handled_checkBox)
        #项目ID（临时生成一个新的ID）

        curr_ID = self.IDget.get()
        ID_text = QTextEdit()
        ID_text.setText(curr_ID)
        ID_text.setReadOnly(True)
        self.tableWidget.setCellWidget(row_count , self.col_name_index[ '项目ID' ] , ID_text)
        #初始化其他单元格
        for header in self.table_header :
            if header in ['机会情况' , '待办程度', '添加追踪', '管线进度', '上线跟进', '明确机会' , '预期订单','待解决', '高亮关注', '项目ID' ] :
                pass
            elif header in [ '子任务']:#如果子任务不存在，则添加追踪和待解决为not enabled
                head_item_textEdit = QTextEdit()
                head_item_textEdit.focusOutEvent = types.MethodType(self.task_edit_focusOutEvent, head_item_textEdit)
                head_item_textEdit.setTabChangesFocus(True)
                self.tableWidget.setCellWidget(row_count , self.col_name_index[ header ] , head_item_textEdit)
            else:
                head_item_textEdit = QTextEdit()
                head_item_textEdit.setTabChangesFocus(True)
                self.tableWidget.setCellWidget(row_count , self.col_name_index[ header ] , head_item_textEdit)
        self.tableWidget.cellWidget(row_count,self.col_name_index['项目名称']).setReadOnly(True)

    def init_new_empty_line(self,proj_name,client_id, row_data:dict = None):
        '''对新的空行填入已有记录数据'''
        # 如果未提供row_data，则根据项目名称和客户名称来获取已有的配置信息，对空行的空间进行参数初始化，同时将这些信息保存在self.log_id_track中
        # 如果提供了row_data, 则直接使用row_data
        row_count = self.tableWidget.rowCount()-1
        #抽取原有记录，用以初始化行内容, 并对已有的空行中自动生成的_id进行覆盖
        if not row_data:
            if proj_name == '' or client_id =='':
                return
            else:pass
        else:
            proj_name = row_data['product']
            client_name = row_data['client']
        #为这一行的log分配一组id
        #即使内容为空，同样也预分配一行id
        row_id_track = {}
        row_id_track['project_name'] = self.lineEdit_2.text()
        row_id_track['task_id'] = self.task_ID_get.get()
        row_id_track['todo_id'] = self.todo_id_get.get()
        row_id_track['memo_id'] = self.memo_id_get.get()
        row_id_track['meeting_id'] = self.meeting_id_get.get()
        row_id_track['status_log_id'] = self.status_id_get.get()#如果该project原先已经存在status_log，则后续用原有的将其覆盖
        row_id_track['is_new_project'] = True
        row_id_track['saved'] = False
        row_id_track[self.col_name_key['机会情况']] = 0
        row_id_track[self.col_name_key['待办程度']] = 0
        row_id_track[self.col_name_key['明确机会']] = False
        row_id_track[self.col_name_key['预期订单']] = False
        row_id_track[self.col_name_key['上线跟进']] = False
        row_id_track[self.col_name_key['待解决']] = 'A'
        row_id_track[self.col_name_key['高亮关注']] = False
        row_id_track[self.col_name_key['添加追踪']] = False
        row_id_track[self.col_name_key['项目名称']] = self.lineEdit_2.text()
        row_id_track[self.col_name_key['客户名称']] = self.client
        row_id_track['client_id'] = self.client_id

        #project从数据库提取该项目原有基本信息
        extract_line = CS.getLinesFromTable(table_name='proj_list', conditions={'product':proj_name, 'client_id':client_id})
        extract_line_keys = extract_line.pop()
        task_desc = ''
        if row_data:#优先采用从会议记录中传入的行信息
            task_desc = row_data['task_desc']
            self.set_value_to_row(row_count, row_data)
            row_id_track['task_id'] = row_data['task_id']
            if 'todo_id' in row_data.keys():#针对旧的版本
                row_id_track['todo_id'] =row_data['todo_id']
            else:
                row_id_track['todo_id'] = self.todo_id_get.get()
            row_id_track['memo_id'] = row_data['memo_id']
            row_id_track['meeting_id'] = row_data['meeting_id']
            row_id_track['status_log_id'] = row_data['status_log_id']
            try:
                row_id_track['is_new_project'] = row_data['is_new_project']
                row_id_track['saved'] = row_data['saved']
            except:
                pass
        elif extract_line:#使用从数据库读取的已有信息
            extract_line_dict = dict(zip(extract_line_keys,extract_line[0]))
            project_id = extract_line_dict['_id']
            task_desc = str(self.textEdit_4.toPlainText()).strip()
            status_id_log = CS.getLinesFromTable('proj_status_log',conditions={'conn_project_id':project_id},
                                                 columns_required=['_id','conn_project_id'])
            status_id_log.pop()

            if status_id_log:
                row_id_track['status_log_id'] = status_id_log[0][0]#如果该project原先已经存在status_log，则用原有的将其覆盖
            #从数据库读取原有的状态信息
            #机会标签
            chance_tag_code = CS.getLinesFromTable('proj_status_log', conditions={'conn_project_id':project_id},
                                                   columns_required=['status_code'])[0][0]
            if not isinstance(chance_tag_code, int):
                chance_tag_code = 0
            if extract_line_dict['current_task_num']:#
                task_info = CS.getLinesFromTable('tasks', conditions={'conn_project_id':project_id,
                                                                      'inter_order_weight':extract_line_dict['current_task_num']},
                                                                      columns_required=['is_critical','officejob_type'])
                task_info.pop()
                if task_info:
                    is_critical = task_info[0][0]
                    officejob_type = task_info[0][1]
                else:
                    is_critical = 0
                    officejob_type = 'A'
            else:
                is_critical = 0
                officejob_type = 'A'
            #此处使用字典较多，字典在类初始化时完成初始化
            #根据原有记录初步填写行内容，完成初始化
            #机会
            time2 = time.time()
            # 如果该客户名下已有该项目信息，则适用已有的信息进行初始化
            # 但如果该项目信息是从会议纪要临时文件中读入的，则接下来用临时文件里的信息来覆盖行控件中的信息，而不覆盖row_id_track中的信息
            row_id_track['is_new_project'] = False

            self.tableWidget.cellWidget(row_count,self.col_name_index['机会情况']).\
                setCurrentIndex(chance_tag_code)
            row_id_track[self.col_name_key['机会情况']] = chance_tag_code
            #程度
            self.tableWidget.cellWidget(row_count,self.col_name_index['待办程度']).\
                setCurrentIndex(is_critical)
            row_id_track[self.col_name_key['待办程度']] = is_critical
            #添加追踪
            self.tableWidget.cellWidget(row_count,self.col_name_index['添加追踪']).setChecked(self.set_todo_checkBox.isChecked())
            #管线进度
            self.tableWidget.cellWidget(row_count,self.col_name_index['管线进度']).setCurrentIndex(self.comboBox_progress.currentIndex())
            row_id_track[self.col_name_key['管线进度']] =DataCenter.progress_code[extract_line_dict[self.col_name_key['管线进度']]]
            #明确机会
            self.tableWidget.cellWidget(row_count,self.col_name_index['明确机会']).\
                setChecked(extract_line_dict[self.col_name_key['明确机会']])
            row_id_track[self.col_name_key['明确机会']] = extract_line_dict[self.col_name_key['明确机会']]
            #预期订单
            self.tableWidget.cellWidget(row_count,self.col_name_index['预期订单']).\
                setChecked(extract_line_dict[self.col_name_key['预期订单']])
            row_id_track[self.col_name_key['预期订单']] = extract_line_dict[self.col_name_key['预期订单']]
            #上线跟进
            self.tableWidget.cellWidget(row_count,self.col_name_index['上线跟进']).\
                setChecked(extract_line_dict[self.col_name_key['上线跟进']])
            row_id_track[self.col_name_key['上线跟进']] = extract_line_dict[self.col_name_key['上线跟进']]
            #待解决
            self.tableWidget.cellWidget(row_count,self.col_name_index['待解决']).\
                setCurrentIndex(list(self.office_job_dict.keys()).index(officejob_type))
            row_id_track[self.col_name_key['待解决']] = officejob_type
            #高亮关注
            self.tableWidget.cellWidget(row_count,self.col_name_index['高亮关注']).\
                setChecked(extract_line_dict[self.col_name_key['高亮关注']])
            row_id_track[self.col_name_key['高亮关注']] = extract_line_dict[self.col_name_key['高亮关注']]
            # #项目备注
            # self.tableWidget.cellWidget(row_count,self.col_name_index['项目备注']).\
            #     setText(extrat_line_dict[self.col_name_key['项目备注']])
            #项目ID
            self.tableWidget.cellWidget(row_count,self.col_name_index['项目ID']).\
                setText(extract_line_dict[self.col_name_key['项目ID']])
            time3 = time.time()
            #print('读数据库用时=',time2-time1,'填充数据用时=',time3-time2)
        else:
            task_desc = str(self.textEdit_4.toPlainText()).strip()
            row_id_track['conn_project_id'] = True#无意义的语句？
            self.tableWidget.cellWidget(row_count,self.col_name_index['添加追踪']).setChecked(self.set_todo_checkBox.isChecked())
            self.tableWidget.cellWidget(row_count,self.col_name_index['管线进度']).setCurrentIndex(self.comboBox_progress.currentIndex())
            pass
        #将为各条记录分配的id进行添加
        row_id_track['conn_project_id'] = self.tableWidget.cellWidget(row_count, self.col_name_index['项目ID']).\
                toPlainText()
        row_id_track['project_name'] = proj_name
        self.log_id_track[row_id_track['conn_project_id']] = row_id_track
        #项目名称
        self.tableWidget.cellWidget(row_count,self.col_name_index['项目名称']).setText(proj_name)
        #客户名称
        self.tableWidget.cellWidget(row_count,self.col_name_index['客户名称']).setText(self.client)
        self.tableWidget.cellWidget(row_count,self.col_name_index['客户名称']).setReadOnly(True)
        #根据子任务设置行控件
        if not task_desc:
            self.tableWidget.cellWidget(row_count,self.col_name_index['待解决']).setEnabled(False)
            self.tableWidget.cellWidget(row_count,self.col_name_index['添加追踪']).setEnabled(False)
            self.tableWidget.cellWidget(row_count,self.col_name_index['添加追踪']).setChecked(False)

    def del_line(self):
        dlt = QMessageBox.warning(self, "注意", "删除可不能恢复了哦！", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if dlt == QMessageBox.No:
            return
        else:
            # curow = self.tableWidget.currentRow()
            selec_model = self.tableWidget.selectionModel()
            selec_rows = selec_model.selectedRows()
            rows = []
            for rr in selec_rows:
                rows.append(rr.row())
            rows.reverse()
            for i in rows:
                conn_project_id = self.tableWidget.cellWidget(i, self.col_name_index['项目ID']).toPlainText()
                project = self.log_id_track.pop(conn_project_id)
                project_name = project['project_name']
                is_new_project = project['is_new_project']
                if is_new_project:#是新的项目，根据名字查找是否已经在删除队列里面
                    for key, value in self.log_id_track.items():
                        if project_name == value['project_name']:
                            break
                    else:
                        self.log_id_track_deleted[conn_project_id] = project
                elif not conn_project_id in self.log_id_track_deleted.keys():#是老项目，根据id查找
                    self.log_id_track_deleted[conn_project_id] = project
                else:
                    pass
                self.tableWidget.removeRow(i)
                #记录被删除的_id
            print(selec_model,'\n',selec_rows, '\n',rows)

    def on_upButton_clicked(self):
        nFrom = self.tableWidget.currentRow()
        if nFrom == 0 or nFrom == -1:
            return
        nTo = nFrom - 1
        self.move_row(nFrom, nTo)

    def on_downButton_clicked(self):
        nFrom = self.tableWidget.currentRow()
        if nFrom == self.tableWidget.rowCount() - 1 or nFrom ==-1:
            return
        nTo = nFrom + 2
        self.move_row(nFrom, nTo)

    def move_row(self, nFrom, nTo):
        row_data = self.get_row_data(nFrom)
        self.add_line_empty( nTo )
        self.set_value_to_row(nTo, row_data)
        if nFrom > nTo:
            nRemove = nFrom + 1
            nCurrent = nTo
        else:
            nRemove = nFrom
            nCurrent = nTo - 1
        self.tableWidget.removeRow(nRemove)
        self.tableWidget.setCurrentCell(nCurrent,0)

    def set_value_to_row(self, nTo, row_data:dict):
        has_task = True if row_data[self.col_name_key['子任务']] else False
        for header in self.table_header:
            if header == '添加追踪':
                if has_task:
                    self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setChecked(
                    row_data[self.col_name_key[header]] )
                else:
                    self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setEnabled(False)
            elif header in ['明确机会', '预期订单', '上线跟进', '高亮关注']:
                self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setChecked(
                    row_data[self.col_name_key[header]] )
            elif header in ['项目名称', '详情记录', '子任务', '项目备注', '客户名称', '项目ID']:
                self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setText(
                    row_data[self.col_name_key[header]] )
            elif header in ['机会情况', '待办程度']:
                self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setCurrentIndex(
                    row_data[self.col_name_key[header]] )
            elif header == '管线进度':
                index = DataCenter.progress_code[row_data[self.col_name_key[header]]]
                self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setCurrentIndex( index )
            elif header == '待解决':
                if has_task:
                    index = 'ABCDE'.index( row_data[self.col_name_key[header]] )
                    self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setCurrentIndex( index )
                else:
                    self.tableWidget.cellWidget( nTo, self.col_name_index[header] ).setEnabled(False)

    def get_table_data(self):
        '''读取tableWidget中的信息'''
        row_count = self.tableWidget.rowCount()
        table_data = []
        for i in range(row_count):
            r = self.get_row_data(i)
            table_data.append(r)
        return table_data

    def get_row_data(self,nRow):
        r = {}
        # '项目名称','详情记录','机会情况','待办程度','子任务','项目备注','明确机会','预期订单','上线跟进','待解决','高亮关注','客户名称','项目ID'
        for header in self.table_header:  #
            if header in ['添加追踪','明确机会','预期订单','上线跟进','高亮关注']:
                r[self.col_name_key[header]] = self.tableWidget.cellWidget( nRow,
                                                                            self.col_name_index[header] ).isChecked()
            elif header in ['项目名称','详情记录','子任务','项目备注','客户名称','项目ID']:
                r[self.col_name_key[header]] = self.tableWidget.cellWidget( nRow,
                                                                            self.col_name_index[header] ).toPlainText()
            else:
                r[self.col_name_key[header]] = self.tableWidget.cellWidget( nRow,
                                                                            self.col_name_index[header] ).currentIndex()
        r['officejob_type'] = 'ABCDE'[r['officejob_type']]
        r['current_progress'] = DataCenter.progress_code_r[r['current_progress']]
        r['task_id'] = self.log_id_track[r[self.col_name_key['项目ID']]]['task_id']
        r['todo_id'] = self.log_id_track[r[self.col_name_key['项目ID']]]['todo_id']
        r['memo_id'] = self.log_id_track[r[self.col_name_key['项目ID']]]['memo_id']
        r['meeting_id'] = self.log_id_track[r[self.col_name_key['项目ID']]]['meeting_id']
        r['status_log_id'] = self.log_id_track[r[self.col_name_key['项目ID']]]['status_log_id']
        r['is_new_project'] = self.log_id_track[r[self.col_name_key['项目ID']]]['is_new_project']
        r['saved'] = self.log_id_track[r[self.col_name_key['项目ID']]]['saved']
        return r

    def get_meeting_meta_info(self, table_data):
        company_id = self.client_id
        company_name = self.client
        date = self.dateEdit.date().toPyDate()
        meeting_date = date.strftime( '%Y-%m-%d' )
        #company_desc = self.textEdit_5.toPlainText()
        # company_meeting_people = self.textEdit.toPlainText()
        # our_people = self.textEdit_2.toPlainText()
        company_meeting_log = {}
        company_meeting_log['project_log'] = table_data
        company_meeting_log['our_people'] = self.textEdit_2.toPlainText()
        company_meeting_log['visited_people'] = json.dumps(self.person_in_meeting,ensure_ascii=False)
        _company_item_log = json.dumps(company_meeting_log,ensure_ascii=False)##会议涉及的实际内容，用字典转换成json存储
        if self.company_meeting_id:
            _meeting_id = self.company_meeting_id
        else:
            self.company_meeting_id = Snow('cmt').get()
            _meeting_id = self.company_meeting_id
        meeting_meta_info = {}
        meeting_meta_info['_id'] = _meeting_id
        meeting_meta_info['conn_company_id'] = self.client_id
        meeting_meta_info['meeting_date'] = date
        company_desc_text = self.textEdit_5.toPlainText()
        company_desc_id = self.client_log_id#company_desc和client_log共用id
        company_desc = {}
        company_desc['id'] = company_desc_id
        company_desc['text'] = company_desc_text
        meeting_meta_info['company_desc'] = json.dumps(company_desc,ensure_ascii=False)
        # meeting_meta_info['company_desc_id'] = self.client_log_id
        meeting_meta_info['meeting_info'] = _company_item_log
        meeting_meta_info['create_time'] = time.strftime("%Y-%m-%d %X", time.localtime())
        return meeting_meta_info

    def save_to_word(self):
        content_list = self.get_table_data()
        document = Document()
        #格式初始化
        document.styles['Normal'].font.name = u'宋体'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        document.styles['Normal'].font.size = Pt(10.5)
        #print('断点3')
        document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        #
        #会议纪要抬头部分
        p = document.add_paragraph(style = 'Normal')
        p.add_run('Dear All，\n\n以下是'+ self.dateEdit.text().split('/')[1] +'月' +
                  self.dateEdit.text().split('/')[2] + '日拜访' + self.lineEdit.text() +'会议纪要：')
        #会议基本信息部分：日期、公司、参会人员
        ##日期
        p = document.add_paragraph(style = 'Normal')
        p.add_run(' ')
        p = document.add_paragraph(style = 'Normal')
        p.add_run(' ')
        p = document.add_paragraph(style = 'Normal')
        p.add_run( self.dateEdit.text().split('/')[0] +'年'+
            self.dateEdit.text().split('/')[1] +'月' +
                  self.dateEdit.text().split('/')[2] + '日').bold=True
        ##客户人员
        p=document.add_paragraph(style = 'Normal')
        p.add_run(self.lineEdit.text()).bold = True
        p = document.add_paragraph(style = 'Normal')
        p.add_run(self.textEdit.toPlainText()).bold=True
        ##安信人员
        p = document.add_paragraph(style = 'Normal')
        p.add_run('广州安信：'+self.textEdit_2.toPlainText())

        #公司情况
        if self.textEdit_5.toPlainText():
            document.add_paragraph(style = 'Normal')
            p = document.add_paragraph(style='Normal')
            p.add_run('公司情况\n').bold = True
            p.add_run(self.textEdit_5.toPlainText())
        #项目列表部分
        index_count = len(content_list)
        for i in range(index_count):
            document.add_paragraph(style = 'Normal')
            h = document.add_heading('',level=1)
            run = h.add_run(str(i+1)+'.'+content_list[i][self.col_name_key['项目名称']])
            run.font.name = u'宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(10.5)
            document.add_paragraph('【管线进度】' + DataCenter.progress_text[content_list[i][self.col_name_key['管线进度']]]
                                   , style = 'Normal')
            # document.add_paragraph(DataCenter.progress_text[content_list[i][self.col_name_key['管线进度']]], style = 'Normal')
            document.add_paragraph('【项目情况】', style = 'Normal')
            document.add_paragraph(content_list[i][self.col_name_key['详情记录']], style = 'Normal')
            if content_list[i][self.col_name_key['子任务']] == '' or content_list[i][self.col_name_key['子任务']] == 'nan':
                pass
            else:
                document.add_paragraph('【需跟进】', style = 'Normal')
                document.add_paragraph(content_list[i][self.col_name_key['子任务']], style = 'Normal')
        for par in document.paragraphs:
            par.paragraph_format.space_before = Pt(0)
            par.paragraph_format.space_after = Pt(0)
        return document

    def save_word_file(self):
        filename = self.dateEdit.date().toString("yyyyMMdd")+ '-' + self.lineEdit.text() + '-'+','.join(self.textEdit_2.toPlainText().split())
        fname, ftype = QFileDialog.getSaveFileName(self,'save file', self.filepath_model.get_last_save()+'\\'+filename , "WordFile (*.docx)")
        print(fname, ftype)
        if fname:
            (filepath, tempfilename) = os.path.split(fname)
            self.filepath_model.save_last_save(filepath)
            doc_body = self.save_to_word()
            try:
                doc_body.save(fname)
            except PermissionError as e:
                print(e)
                QMessageBox.about(self, '保存失败', '检查文件是否被其他程序打开')
        else:
            return

    def save_to_excel(self):
        content_list = self.get_table_data()
        return content_list

    def save_excel_file(self):
        filename =  self.dateEdit.date().toString("yyyyMMdd") + '-' + self.lineEdit.text() + '-'+','.join(self.textEdit_2.toPlainText().split())
        fname, ftype = QFileDialog.getSaveFileName(self, 'save file',
                                                   self.filepath_model.get_last_save() + '\\' + filename ,
                                                   "ExcelFile (*.xlsx)")
        print(fname, ftype)
        if fname:
            #保存本次选择的路径
            (filepath, tempfilename) = os.path.split(fname)
            self.filepath_model.save_last_save(filepath)
            #保存xlsx文件

            '''work_book = xlwt.Workbook()
            work_sheet = work_book.add_sheet(sheetname='Sheet1')
            sheet_body = self.save_to_excel()
            for i, field in enumerate(self.table_header):
                work_sheet.write(0,i,field)
            for i, row in enumerate(sheet_body):
                for j, field in enumerate(self.table_header):
                    work_sheet.write(i+1,j, row[self.col_name_key[field]])
            work_book.save(fname)'''
            work_book = openpyxl.Workbook()
            work_sheet = work_book.active
            sheet_body = self.get_table_data()
            company_id = self.client_id
            company_name = self.client
            date = self.dateEdit.date().toPyDate()
            meeting_date = date.strftime('%Y-%m-%d')
            company_desc = self.textEdit_5.toPlainText()
            company_meeting_people = self.textEdit.toPlainText()
            our_people = json.dumps(self.person_in_meeting,ensure_ascii=False)
            work_sheet.append([company_name, meeting_date, company_desc, company_meeting_people, our_people,company_id])
            work_sheet.append(self.table_header)
            for i, row in enumerate(sheet_body):
                for j, field in enumerate(self.table_header):
                    work_sheet.cell(i+3,j+1, value=row[self.col_name_key[field]])#openpyxl从1开始计数，并且默认没有表头
            try:
                work_book.save(fname)
            except PermissionError as e:
                print(e)
                QMessageBox.about(self, '保存失败', '检查文件是否被其他程序打开')
                return
        else:
            return

    def load_from_excel(self):
        '''这个地方有问题，没有考虑到log_id_track'''
        fname, ftype = QFileDialog.getOpenFileName(self, 'open file', self.filepath_model.get_last_open(),
                                                        "ExcelFile (*.xlsx)")
        if fname:

            (filepath, tempfilename) = os.path.split(fname)
            self.filepath_model.save_last_open(filepath)
            print(fname)
            work_book = openpyxl.load_workbook(fname, data_only= True)
            excel_data = work_book.active
        else:
            return
        #print(excel_data)
        meeting_meta_info = next(excel_data.values)
        excel_header = next(excel_data.values)

        #company_name, meeting_date, company_desc, company_meeting_people, our_people, company_id
        self.client_id = meeting_meta_info[5]
        client_name_get = CS.getLinesFromTable('clients', {'_id':self.client_id}, ['short_name'])
        client_name_get.pop()
        self.client = client_name_get[0][0]
        # try:
        #     self.client = next(excel_data.values)[self.col_name_index[ '客户名称']]
        # except :
        #     QMessageBox.about(self,'无记录', '文件中无记录')
        #     return
        self.lineEdit.setText(self.client)
        if isinstance(meeting_meta_info[1],str):
            date = time.strptime(meeting_meta_info[1],'%Y-%m-%d')
        elif isinstance(meeting_meta_info[1], datetime.datetime):
            date = meeting_meta_info[1].timetuple()
        else:
            QMessageBox.about(self, '错误', '文件内容格式错误')
            return
        self.dateEdit.setDateTime(QDateTime(QDate(date.tm_year,date.tm_mon, date.tm_mday)))
        self.textEdit_5.setText(meeting_meta_info[2])
        # self.textEdit.setText(meeting_meta_info[3])#客户人员
        self.init_comboBox_personnel_checked(meeting_meta_info[3])
        self.textEdit_2.setText(meeting_meta_info[4])#我方人员
        excel_data.delete_rows(1,2)#删除表头和元信息
        self.log_id_track.clear()
        self.tableWidget.setRowCount(0)
        try:
            for i ,row in enumerate(excel_data.values):
                #根据当前所在的导入数据行，在项目动态数据文件里匹配得出extract_line
                proj_name = row[self.col_name_index['项目名称']]
                client_name = row[self.col_name_index[ '客户名称']]
                time1 = time.time()
                self.add_line_empty()#增加新的空行，注意确保i和该函数内部的row_count变量应该相等
                time2 = time.time()
                #project_id始终遵循原有的生成规则，不从Excel导入
                #
                self.init_new_empty_line(proj_name, client_name)
                time3 = time.time()
                #详情记录
                print(self.tableWidget.rowCount())
                self.tableWidget.cellWidget(i,self.col_name_index['详情记录']).setText(str(row[ self.col_name_index['详情记录']]))
                #子任务
                if row[self.col_name_index['子任务']]:
                    self.tableWidget.cellWidget(i,self.col_name_index['子任务']).setText(str(row[self.col_name_index['子任务']]))
                #项目备注
                if row[self.col_name_index['项目备注']]:
                    self.tableWidget.cellWidget(i,self.col_name_index['项目备注']).setText(str(row[self.col_name_index['项目备注']]))
                #
                self.tableWidget.cellWidget(i, self.col_name_index['机会情况']).\
                    setCurrentIndex(row[self.col_name_index['机会情况']])
                #程度
                self.tableWidget.cellWidget(i,self.col_name_index['待办程度']).\
                    setCurrentIndex(row[self.col_name_index['待办程度']])
                #明确机会
                self.tableWidget.cellWidget(i,self.col_name_index['明确机会']).\
                    setChecked(row[self.col_name_index['明确机会']])
                #预期订单
                self.tableWidget.cellWidget(i,self.col_name_index['预期订单']).\
                    setChecked(row[self.col_name_index['预期订单']])
                #上线跟进
                self.tableWidget.cellWidget(i,self.col_name_index['上线跟进']).\
                    setChecked(row[self.col_name_index['上线跟进']])
                #待解决
                self.tableWidget.cellWidget(i,self.col_name_index['待解决']).\
                    setCurrentIndex(list(self.office_job_dict.keys()).index(row[self.col_name_index['待解决']]))
                #高亮关注
                self.tableWidget.cellWidget(i,self.col_name_index['高亮关注']).\
                    setChecked(row[self.col_name_index['高亮关注']])
                time4 = time.time()
                print('添加空行用时=', time2-time1,'初始化空行用时=',time3-time2, '填充新数据用时=',time4-time3)
        except:
            QMessageBox.about(self, '错误', '文件内容格式错误')
            return
        self.set_project_completer()

    def load_from_meeting_log(self):
        company_meeting = CS.getLinesFromTable('client_meeting_log',
                                                conditions={'_id':self.company_meeting_id})
        fields = company_meeting.pop()
        meeting_factors = dict(zip(fields,company_meeting[0]))
        date = time.strptime(meeting_factors['meeting_date'],'%Y-%m-%d')
        meeting_info = json.loads(meeting_factors['meeting_info'])
        self.dateEdit.setDateTime(QDateTime(QDate(date.tm_year,date.tm_mon, date.tm_mday)))
        try:
            company_desc = json.loads(meeting_factors['company_desc'])
            company_desc_text = company_desc['text']
            company_desc_id = company_desc['id']
            self.textEdit_5.setText(company_desc_text)
            self.client_log_id = company_desc_id
        except:
            self.textEdit_5.setText(meeting_factors['company_desc'])
        #客户人员
        visited_people = meeting_info['visited_people']
        self.init_comboBox_personnel_checked(visited_people)
        self.textEdit_2.setText(meeting_info['our_people'])
        self.lineEdit.setText(self.client)
        if meeting_info['project_log'] and 'status_log_id' in meeting_info['project_log'][0].keys():#按照新的方式存储的会包含_id字段
            pass
        elif isinstance(meeting_info['project_log'] ,list) and not meeting_info['project_log']:
            pass#project_log 是空列表
        else:
            self.close()
            QMessageBox.about(self,'错误', '会议记录无法读取')

            return
        for i, row in enumerate(meeting_info['project_log']):
            product = row['product']
            if 'current_progress' not in row.keys():#针对旧的json存储内容
                row['current_progress'] = 'inv'
            if 'todo' not in row.keys():#针对旧的json存储内容
                row['todo'] = False
            client = row['client']
            self.add_line_empty()
            self.init_new_empty_line(product, client, row)

    def closeEvent(self, event):
        reply = QMessageBox.question(self, '警告', '退出后会议内容不可再修改！\n确认退出吗？',
                                           QMessageBox.Yes| QMessageBox.No, QMessageBox.No)
        if reply == QtWidgets.QMessageBox.Yes:
            event.accept()
        else:
            event.ignore()

    def getProjectWeight(self, project, status_code):
        '''计算project的weight的值'''
        weight = 0
        # if project.is_critical:
        #     weight += 3
        if project['in_act']:
            weight += 2
        if project['order_tobe']:
            weight += 3
        if project['clear_chance']:
            weight += 3
        if project['current_progress'] == 'ini' :
            weight += 2
        elif project['current_progress'] == 'inv':
            weight += 1
        if DataCenter.sequence_chance_code_dict[status_code] == 1:
            weight *= 1.2
        elif DataCenter.sequence_chance_code_dict[status_code] >= 4:
            weight*=0.8
        if project['highlight']:
            weight *= 1.5
        return weight
    # def on_edit_again(self):
    #     message = QMessageBox(self)
    #     ok = message.question(self,'再次编辑', '是否继续会议的信息？\n请确认已完成编辑。')
    #     if ok == QMessageBox.No:
    #         return
    #     self.secondary_editing = True
    #     self.pushButton.setEnabled(True)
    #     self.pushButton_9.setEnabled(False)
    #     pass

    def update_dynamic(self):
        '''将会议信息更新到后台和整个UI页面'''
        table_data = self.get_table_data()
        if len(table_data) == 0 and self.textEdit_5.toPlainText().strip() == '' and not self.log_id_track_deleted:
            QMessageBox.warning(self, '未输入', '请输入会议记录！')
            return
        if self.textEdit.toPlainText().strip() == '' or self.textEdit_2.toPlainText().strip() == '':
            QMessageBox.warning(self, '人员信息', '请完善参会人员！')
            return
        message = QMessageBox(self)
        ok = message.question(self,'保存', '是否保存此次会议的信息？\n请确认已完成编辑。')
        if ok == QMessageBox.No:
            return
        self.update_new_log(table_data)
        self.save_meeting_info(table_data)
        self.reverse_deleted_log()
        QMessageBox.about(self,'已保存', '会议纪要已保存。')
        self.parent.listener.rerender()

    def update_new_log(self, table_data):
        '''保存每一条项目的更新信息到项目记录，保存公司信息（如果有）并使用提前分配到的ID'''
        IDs = [row['_id'] for row in table_data]
        # print(IDs)
        ID_len = len(IDs)
        project_fields = ['product', 'current_progress','clear_chance', 'order_tobe', 'in_act', 'highlight', 'client', '_id']
        task_fields = ['is_critical', 'task_desc', 'officejob_type', 'inter_order_weight', '_id', 'conn_project_id']
        memo_log_fields = ['memo_desc', '_id', 'create_time', 'inter_order_weight', 'conn_project_id']
        meeting_log_fields = ['meeting_desc', '_id', 'create_time', 'inter_order_weight', 'conn_project_id']
        proj_status_log_fields = ['status_code', '_id', 'conn_project_id']
        company_project_log = []
        for i, row in enumerate(table_data):
            #首先获取各个表对应的字段
            #project字段
            project_fields_values = {}
            for field in project_fields:
                project_fields_values[field] = row[field]

            project_fields_values['client_id'] = self.client_id
            conn_project_id = project_fields_values['_id']
            #weight计算
            project_fields_values['weight'] = self.getProjectWeight(project_fields_values, row['status_code'] )
            self.log_id_track[conn_project_id]['saved'] = True#该项目已经被保存过
            #task字段
            task_fields_values = {}
            todo_fields_values = {}
            if row['task_desc'].strip() == '' or not row['todo']:
                pass
            else:
                #task
                for field in task_fields:
                    try:
                        task_fields_values[field] = row[field]
                    except:
                        continue
                #此处如果有如果有task,且项目为新项目，则project的current_task_num变为1
                if self.log_id_track[conn_project_id]['is_new_project']:
                    project_fields_values['current_task_num'] = 1
                task_fields_values['inter_order_weight'] = len(CS.getLinesFromTable('tasks',conditions={
                    'conn_project_id':row['_id']},columns_required=['inter_order_weight','_id'])) -1
                task_fields_values['conn_project_id'] = row['_id']
                task_fields_values['_id'] = self.log_id_track[conn_project_id]['task_id']#从分配好的id中读取
                task_fields_values['create_time'] = time.strftime("%Y-%m-%d %X", time.localtime())
                #todo_
                todo_fields_values['_id'] = self.log_id_track[conn_project_id]['todo_id']
                todo_fields_values['conn_task_id'] = task_fields_values['_id']
                todo_fields_values['conn_project_id'] = task_fields_values['conn_project_id']
                todo_fields_values['todo_desc'] = task_fields_values['task_desc']
                todo_fields_values['status'] = 0
                todo_fields_values['is_critical'] = task_fields_values['is_critical']
                todo_fields_values['create_time'] = task_fields_values['create_time']
                todo_fields_values['officejob_type'] = task_fields_values['officejob_type']

            #memo字段
            memo_log_fields_values = {}
            if row['memo_desc'].strip() == '':
                pass
            else:
                for field in memo_log_fields:
                    try:
                        memo_log_fields_values[field] = row[field]#这里得到的_id是错误的，在后面被memo_id覆盖
                    except:
                        continue
                memo_log_fields_values['inter_order_weight'] = len(CS.getLinesFromTable('project_memo_log',conditions={
                    'conn_project_id':row['_id']},columns_required=['inter_order_weight','_id'])) -1
                memo_log_fields_values['conn_project_id'] = row['_id']
                memo_log_fields_values['_id'] = self.log_id_track[conn_project_id]['memo_id']#
                memo_log_fields_values['create_time'] = time.strftime("%Y-%m-%d %X", time.localtime())
            #meeting字段
            meeting_log_fields_values = {}
            for field in meeting_log_fields:
                try:
                    meeting_log_fields_values[field] = row[field]
                except:
                    continue
            meeting_log_fields_values['inter_order_weight'] = len(CS.getLinesFromTable('project_meeting_log',conditions={
                'conn_project_id':row['_id']},columns_required=['inter_order_weight','_id'])) -1
            meeting_log_fields_values['conn_project_id'] = row['_id']
            meeting_log_fields_values['_id'] = self.log_id_track[conn_project_id]['meeting_id']
            meeting_log_fields_values['create_time'] = time.strftime("%Y-%m-%d %X", time.localtime())
            #status字段
            proj_status_log_fields_values = {}
            for field in proj_status_log_fields:
                try:
                    proj_status_log_fields_values[field] = row[field]
                except:
                    continue
            proj_status_log_fields_values['conn_project_id'] = row['_id']
            proj_status_log_fields_values['_id'] = self.log_id_track[conn_project_id]['status_log_id']
            # proj_status_log_fields_values['_id'] = self.task_ID_get.get()
            # proj_status_log_fields_values['create_time'] = time.strftime("%Y-%m-%d %X", time.localtime())
            #当前保留信息写入数据库
            CS.upsertSqlite('proj_list',keys=list(project_fields_values.keys()),
                            values=list(project_fields_values.values()))
            if task_fields_values:
                CS.upsertSqlite('tasks', keys=list(task_fields_values.keys()), values=list(task_fields_values.values()))
                CS.upsertSqlite('todo_log', keys=list(todo_fields_values.keys()), values=list(todo_fields_values.values()))
            CS.upsertSqlite('project_meeting_log', keys=list(meeting_log_fields_values.keys()),
                            values= list(meeting_log_fields_values.values()))
            if memo_log_fields_values:
                CS.upsertSqlite('project_memo_log', keys=list(memo_log_fields_values.keys()),
                              values=list(memo_log_fields_values.values()))

            CS.upsertSqlite('proj_status_log', keys=list(proj_status_log_fields_values.keys()),
                            values=list(proj_status_log_fields_values.values()))

        #单独保存一条记录到client_log表
        if not self.textEdit_5.toPlainText().strip() == '':
            client_log = {}
            client_log['_id'] = self.client_log_id
            client_log['company_id'] = self.client_id
            client_log['log_desc'] = self.textEdit_5.toPlainText()#公司情况
            client_log['create_time'] = time.strftime("%Y-%m-%d %X", time.localtime())
            success1 = CS.upsertSqlite('client_log', keys= list(client_log.keys()),
                                       values=list(client_log.values()))

        # if not success1:
        #     # message.about(self,'失败', '保存失败！')
        #     return False
        # else:
        #     return True
        # message.about(self,'完成', '保存成功！')

    def auto_save_meeting_info(self):
        company_desc = self.textEdit_5.toPlainText()
        if self.tableWidget.rowCount() == 0 and company_desc.strip() == '':
            return
        table_data = self.get_table_data()
        self.save_meeting_info(table_data)

    def save_meeting_info(self, table_data):
        meeting_meta_data = self.get_meeting_meta_info(table_data)
        success = CS.upsertSqlite('client_meeting_log', keys= list(meeting_meta_data.keys()),
                                  values= list(meeting_meta_data.values()))
        return success

    def reverse_deleted_log(self):
        # 目的：在一次会议纪要编写窗口关闭前，所有的保存操作均可撤销，并以关闭窗口前最后一次保存的结果为最终结果
        # 初始化：若为已有的project，加载project原有记录中的参数，is_new_project为False；或初始化新参数，并标记为is_new_project为True；
        # 首次保存：全部标记为已保存；
        # 删除：将初始化的参数加入到已删除队列；若is_new_project，则需要先检查项目名称是否已在被删除的队列中；否则根据project_id来检查
        # 删除后重新保存：用track中的id删除追加型记录，或者覆盖覆写类的的记录；其中proj_list和proj_status_log通常为覆写型，
        #             但若is_new_project为True则视为追加型记录
        #
        for key, value in self.log_id_track_deleted.items():
            #project
            if value['saved'] == False:#如果该项目的信息尚未被保存过，则直接忽略
                continue
            project_keys_values = {}
            project_keys_values['_id'] = key
            project_keys_values['client_id'] = value['client_id']
            for header in ['明确机会','预期订单','上线跟进','高亮关注','客户名称','项目名称']:
                project_keys_values[self.col_name_key[header]] = value[self.col_name_key[header]]
            if value['is_new_project']:
                CS.deleteLineFromTable('proj_list',conditions={'_id':key})
            else:
                CS.upsertSqlite('proj_list',keys=list(project_keys_values.keys()),values= list(project_keys_values.values()))
            #task
            task = {}
            task['_id'] = value['task_id']
            #todo_适合task绑定的，
            todo={}
            todo['conn_task_id'] = value['task_id']
            CS.deleteLineFromTable('tasks', conditions=task)
            CS.deleteLineFromTable('todo_log', conditions=todo)
            #memo
            memo = {}
            memo['_id'] = value['memo_id']
            CS.deleteLineFromTable('project_memo_log',conditions=memo)
            #meeting
            meeting = {}
            meeting['_id'] = value['meeting_id']
            CS.deleteLineFromTable('project_meeting_log',conditions=meeting)
            #status_log
            status = {}
            status['_id'] = value['status_log_id']
            status['conn_project_id'] = key
            status['status_code'] = value['status_code']
            CS.updateSqliteCells('proj_status_log',conditions={'conn_project_id':status['conn_project_id']},
                                 update_fields=status)
            #todo:client_log

            #client_meeting_log
        self.log_id_track_deleted.clear()

    def on_editing_personnel(self, data):
        '''编辑人员信息，data是字典组成的列表'''
        personnel_data = []
        personnel_ids = []
        fields = ['name', 'tittle','telephone']
        for row in self.personnel_data:
            tempList = []
            for key in fields:
                tempList.append(row[key])
            personnel_data.append(tempList)
            personnel_ids.append(row['_id'])
        personnel_table = RedefinedWidget.VectorEditTable(self,['姓名','职位','电话'], personnel_data, ['姓名','职位','电话'],
                                          self.pushButton_personnel,column_widths=[58,65,80],old_data_editable=False)
        #personnel_table不是单例，每次退出即销毁
        personnel_table.EditingFinished.connect(lambda new_personnel_data: self.on_personnel_edited(new_personnel_data))
        personnel_table.show()

    def on_personnel_edited(self, new_personnel_data):
        fields = ['name', 'tittle','telephone']
        old_personnel_data = copy.deepcopy(self.personnel_data)
        len_old = len(old_personnel_data)
        psn_id = Snow('psn')
        psl_id = Snow('psl')
        old_names = []
        for i, row in enumerate(new_personnel_data):
            if i < len_old:
                old_names.append(old_personnel_data[i]['name'])
            if i >= len_old:
                if row[0] in old_names:
                    QMessageBox.about(self,'重复添加','%s已存在，将不添加！'%row[0])
                    continue
                tmp_id = psn_id.get()#临时_id
                #查找已有人名
                find_person = self.checkNewPerson(tmp_id, row[0], row[2])
                person_log = PersonLog()
                if find_person:#如果查找到了现有的人名
                    person_id = self.handleNewPerson(find_person,row[0])
                    if person_id is None:#用户取消操作，跳过该人员
                        continue
                    elif person_id:
                        get_id = person_id
                    else:
                        get_id = tmp_id
                else:
                    get_id = tmp_id
                #获取到了id,继续执行
                person_log.conn_person_id = get_id
                person_log._id = psl_id.get()
                person_log.conn_company_id = self.client_id
                person_log.conn_company_name = self.client
                person_log.job_title = row[1]
                person_log.name = row[0]
                person_log.in_service = True
                person_log.addPersonLog()

                old_personnel_data.append({})
                old_personnel_data[i]['_id'] = get_id
                old_personnel_data[i]['email'] = None
                for j, key in enumerate(fields):
                    old_personnel_data[i][key] = row[j]

        new_data = old_personnel_data
        new_personnel_json = json.dumps(new_data,ensure_ascii=False)

        self.personnel_data = copy.deepcopy(new_data)
        CS.upsertSqlite('clients', ['_id', 'short_name', 'personnel'], [self.client_id,self.client, new_personnel_json])
        cmd = DataCenter.GPersonnelCmd('update', self.client_id, source_widget=self)
        self.parent.listener.accept(cmd)
        self.set_comboBox_personnel(is_renew = True)

    def checkNewPerson(self,_id, name, telephone):
        person_get = Person.findPerson(name)
        if not person_get:
            new_person = Person()
            new_person._id = _id
            new_person.name = name
            new_person.telephone = telephone
            new_person.addPerson()
            return None
        else:
            return person_get

    def handleNewPerson(self,find_person, name):
        name_with_log = ['以下都不是']
        ids = []
        for person_info in find_person:
            find_person_log = PersonLog.findPersonLog( person_info['_id'] )  # 根据person_id查询记录
            if find_person_log:
                job_title = find_person_log[-1]['job_title'] if find_person_log[-1]['job_title'] else ''
                log_text = name + '--' + find_person_log[-1]['conn_company_name'] + '--' + job_title
                name_with_log.append( log_text )
                ids.append( person_info['_id'] )
            else:
                continue
        index,ok = ComboBoxDialog.getIndex( self,'重名人员','发现同名记录，\n请确认是否为以下人员：',
                                            name_with_log )
        if not ok:
            return None
        if index == 0 or index == -1:
            return []
        else:
            return ids[index - 1]

    # @pyqtSlot(tuple)
    def on_person_in_meeting(self, check_status):
        self.person_in_meeting = []
        person_in_meeting = []
        for i, checked in enumerate(check_status):
            if checked:
                self.person_in_meeting.append(self.personnel_data[i])
                title = self.personnel_data[i]['tittle']
                if not title:
                    title = ''
                person_in_meeting.append(self.personnel_data[i]['name'] + ' ' + title)
        person_in_meeting_text = '\n'.join(person_in_meeting)
        self.textEdit.setText(person_in_meeting_text)

    def Accept(self,cmd):
        if cmd.source_widget is self:
            return
        if cmd.flag == 5 and cmd._id == self.client_id:
            self.initPersonnelSet()

    def initAllData(self):
        pass

    def closeEvent(self,e):
        self.timer.stop()
        self.auto_save_meeting_info()
        for i,record in enumerate(self.parent.meetModeAdded):
            if self.client_id == record[0]:
                self.parent.meetModeAdded.pop(i)
        self.parent.listener.removeObserver(self)
        if self.parent.isHidden():
            self.parent.show()
        e.accept()
if __name__ == '__main__' :
        work_book = openpyxl.load_workbook('.\\tag_projection.xlsx',)
        work_sheet = work_book.active
        list = work_sheet['C']
